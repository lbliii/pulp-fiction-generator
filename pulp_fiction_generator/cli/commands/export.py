"""
Export commands for Pulp Fiction Generator.
"""

import os
import typer
import shutil
import markdown
from typing import List, Optional
from enum import Enum
from pathlib import Path
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn

from ..base import BaseCommand
from ...utils.story_persistence import StoryPersistence
from ...utils.errors import logger

# Create a Typer app for export commands
export_app = typer.Typer(help="Export stories in various formats")
console = Console()

class OutputFormat(str, Enum):
    """Output formats for story export."""
    plain = "plain"
    markdown = "md"
    html = "html"
    pdf = "pdf"
    docx = "docx"
    epub = "epub"
    all = "all"

def get_exporter_for_format(output_format: str):
    """Get the appropriate exporter function for a format."""
    exporters = {
        "plain": export_plain,
        "md": export_markdown,
        "markdown": export_markdown,
        "html": export_html,
        "pdf": export_pdf,
        "docx": export_docx,
        "epub": export_epub,
    }
    return exporters.get(output_format.lower())


def export_plain(content: str, output_path: str) -> str:
    """Export content as plain text."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def export_markdown(content: str, output_path: str) -> str:
    """Export content as markdown."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def export_html(content: str, output_path: str) -> str:
    """Export content as HTML."""
    try:
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Generated Story</title>
    <style>
        body {{
            font-family: Georgia, serif;
            line-height: 1.6;
            margin: 0 auto;
            max-width: 800px;
            padding: 2rem;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            font-family: 'Helvetica Neue', Arial, sans-serif;
            margin-top: 2rem;
        }}
        h1 {{
            text-align: center;
            margin-bottom: 2rem;
            font-size: 2.5rem;
            border-bottom: 2px solid #333;
            padding-bottom: 0.5rem;
        }}
        p {{
            margin-bottom: 1.2rem;
            text-align: justify;
        }}
        blockquote {{
            font-style: italic;
            border-left: 4px solid #ccc;
            padding-left: 1rem;
            margin-left: 0;
        }}
        @media (max-width: 600px) {{
            body {{
                padding: 1rem;
            }}
            h1 {{
                font-size: 2rem;
            }}
        }}
    </style>
</head>
<body>
    {markdown.markdown(content)}
</body>
</html>
        """
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        return output_path
    except Exception as e:
        logger.error(f"Error exporting to HTML: {e}")
        raise


def export_pdf(content: str, output_path: str) -> str:
    """Export content as PDF."""
    try:
        # First convert to HTML
        html_path = output_path.replace(".pdf", ".html")
        export_html(content, html_path)
        
        # Then use weasyprint to convert HTML to PDF
        try:
            from weasyprint import HTML
            HTML(html_path).write_pdf(output_path)
            
            # Remove temporary HTML file
            os.remove(html_path)
            return output_path
        except ImportError:
            console.print("[yellow]WeasyPrint not installed. Falling back to alternative PDF export.[/yellow]")
            
            # Try alternative PDF export with pdfkit
            try:
                import pdfkit
                pdfkit.from_file(html_path, output_path)
                
                # Remove temporary HTML file
                os.remove(html_path)
                return output_path
            except ImportError:
                console.print("[yellow]pdfkit not installed. Cannot export to PDF.[/yellow]")
                console.print("[yellow]Install with: pip install pdfkit[/yellow]")
                raise ImportError("No PDF export library available")
    except Exception as e:
        logger.error(f"Error exporting to PDF: {e}")
        raise


def export_docx(content: str, output_path: str) -> str:
    """Export content as DOCX."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        document = Document()
        
        # Extract title from the content (first line)
        lines = content.strip().split("\n")
        title = lines[0].replace("#", "").strip() if lines and lines[0].startswith("#") else "Generated Story"
        
        # Add title
        document.add_heading(title, 0)
        
        # Process content by paragraphs
        current_para = ""
        in_heading = False
        
        for line in lines[1:]:
            # Skip the title line
            if line.strip() == "" and current_para:
                # End of paragraph
                document.add_paragraph(current_para)
                current_para = ""
            elif line.startswith("#"):
                # This is a heading
                if current_para:
                    document.add_paragraph(current_para)
                    current_para = ""
                
                # Add the heading (count # to determine level)
                level = 0
                heading_text = line.strip()
                while heading_text.startswith("#"):
                    level += 1
                    heading_text = heading_text[1:]
                
                # Cap at level 9 (docx only supports up to 9)
                level = min(level, 9)
                document.add_heading(heading_text.strip(), level)
            else:
                # Regular paragraph text
                if current_para:
                    current_para += " " + line.strip()
                else:
                    current_para = line.strip()
        
        # Add the last paragraph if any
        if current_para:
            document.add_paragraph(current_para)
        
        # Save the document
        document.save(output_path)
        return output_path
    except ImportError:
        console.print("[yellow]python-docx not installed. Cannot export to DOCX.[/yellow]")
        console.print("[yellow]Install with: pip install python-docx[/yellow]")
        raise ImportError("python-docx not available")
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}")
        raise


def export_epub(content: str, output_path: str) -> str:
    """Export content as EPUB."""
    try:
        from ebooklib import epub
        
        # Extract title from the content (first line)
        lines = content.strip().split("\n")
        title = lines[0].replace("#", "").strip() if lines and lines[0].startswith("#") else "Generated Story"
        
        # Create new EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_title(title)
        book.set_language('en')
        book.add_author('Pulp Fiction Generator')
        
        # Create chapter
        chapter = epub.EpubHtml(title=title, file_name='story.xhtml')
        chapter.content = markdown.markdown(content)
        
        # Add chapter to book
        book.add_item(chapter)
        
        # Create TOC
        book.toc = (epub.Link('story.xhtml', title, title),)
        
        # Add default NCX and Nav files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS
        style = '''
            body {
                font-family: Georgia, serif;
                line-height: 1.6;
                margin: 1em;
                color: #333;
            }
            h1, h2, h3, h4, h5, h6 {
                font-family: Arial, sans-serif;
                margin-top: 1.5em;
            }
            h1 {
                text-align: center;
                margin-bottom: 1.5em;
                font-size: 2em;
                border-bottom: 1px solid #333;
                padding-bottom: 0.5em;
            }
            p {
                margin-bottom: 1em;
                text-align: justify;
            }
            blockquote {
                font-style: italic;
                border-left: 2px solid #ccc;
                padding-left: 1em;
                margin-left: 0;
            }
        '''
        
        # Add CSS file
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Create spine
        book.spine = ['nav', chapter]
        
        # Write to file
        epub.write_epub(output_path, book, {})
        return output_path
    except ImportError:
        console.print("[yellow]EbookLib not installed. Cannot export to EPUB.[/yellow]")
        console.print("[yellow]Install with: pip install ebooklib markdown[/yellow]")
        raise ImportError("EbookLib not available")
    except Exception as e:
        logger.error(f"Error exporting to EPUB: {e}")
        raise


@export_app.command("story")
def export_story_cmd(
    project: str = typer.Argument(..., help="Project name to export"),
    format: List[OutputFormat] = typer.Option(
        [OutputFormat.markdown], "--format", "-f", 
        help="Output format(s) to export to"
    ),
    output_dir: Optional[str] = typer.Option(
        None, "--output", "-o", 
        help="Output directory (defaults to ./exports/PROJECT_NAME)"
    ),
    include_metadata: bool = typer.Option(
        False, "--metadata", "-m",
        help="Include story metadata in exported files"
    ),
):
    """Export a story to various formats."""
    # Initialize story persistence to access stories
    story_dir = os.getenv("OUTPUT_DIR", "./output")
    story_persistence = StoryPersistence(story_dir)
    
    # Set up output directory
    if not output_dir:
        output_dir = f"./exports/{project}"
    
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Load the story
    try:
        story = story_persistence.load_story(project)
        console.print(f"[green]Loaded project: {project}[/green]")
    except Exception as e:
        console.print(f"[bold red]Error loading project '{project}': {str(e)}[/bold red]")
        return
    
    # Get content
    content = story.content
    
    # Add metadata if requested
    if include_metadata:
        metadata = story.metadata
        metadata_text = f"""
# Story Metadata

- **Title**: {metadata.title}
- **Genre**: {metadata.genre}
- **Chapters**: {metadata.chapter_count}
- **Word Count**: {metadata.word_count:,}
- **Generated with**: {metadata.model if hasattr(metadata, "model") else "Unknown model"}
"""
        content = metadata_text + "\n\n" + content
    
    # Process all formats
    formats_to_export = []
    if OutputFormat.all in format:
        formats_to_export = [f for f in OutputFormat if f != OutputFormat.all]
    else:
        formats_to_export = format
    
    # Export to each format
    with Progress() as progress:
        task = progress.add_task("[green]Exporting...", total=len(formats_to_export))
        
        for fmt in formats_to_export:
            try:
                progress.update(task, description=f"[green]Exporting to {fmt}...")
                output_file = output_path / f"{project}.{fmt}"
                
                # Get the exporter function
                exporter = get_exporter_for_format(fmt)
                if not exporter:
                    console.print(f"[yellow]No exporter available for format: {fmt}[/yellow]")
                    continue
                
                # Export the content
                result_path = exporter(content, str(output_file))
                console.print(f"[green]Exported to {fmt}: {result_path}[/green]")
                
            except ImportError as e:
                console.print(f"[yellow]Missing dependency for {fmt} format: {e}[/yellow]")
                continue
            except Exception as e:
                console.print(f"[bold red]Error exporting to {fmt}: {str(e)}[/bold red]")
                continue
            finally:
                progress.update(task, advance=1)
    
    console.print(f"[bold green]Export complete![/bold green] Files saved to: {output_path}")


@export_app.command("bulk")
def bulk_export_cmd(
    format: List[OutputFormat] = typer.Option(
        [OutputFormat.markdown], "--format", "-f", 
        help="Output format(s) to export to"
    ),
    output_dir: str = typer.Option(
        "./exports", "--output", "-o", 
        help="Output directory"
    ),
    include_metadata: bool = typer.Option(
        False, "--metadata", "-m",
        help="Include story metadata in exported files"
    ),
    recent: Optional[int] = typer.Option(
        None, "--recent", "-r",
        help="Export only the most recent N projects"
    ),
):
    """Bulk export multiple stories to various formats."""
    # Initialize story persistence to access stories
    story_dir = os.getenv("OUTPUT_DIR", "./output")
    story_persistence = StoryPersistence(story_dir)
    
    # Set up output directory
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Get all projects
    try:
        projects = story_persistence.list_stories()
        if not projects:
            console.print("[yellow]No projects found.[/yellow]")
            return
            
        # Sort by modification time if recent flag is set
        if recent is not None:
            # Get metadata for all projects to sort by date
            project_data = []
            for proj in projects:
                try:
                    story = story_persistence.load_story(proj)
                    if hasattr(story.metadata, "last_modified"):
                        project_data.append((proj, story.metadata.last_modified))
                    else:
                        project_data.append((proj, ""))
                except Exception:
                    project_data.append((proj, ""))
            
            # Sort by last_modified (most recent first)
            project_data.sort(key=lambda x: x[1], reverse=True)
            
            # Limit to the most recent N projects
            projects = [p[0] for p in project_data[:recent]]
        
        console.print(f"[green]Found {len(projects)} projects to export[/green]")
    except Exception as e:
        console.print(f"[bold red]Error listing projects: {str(e)}[/bold red]")
        return
    
    # Process all formats
    formats_to_export = []
    if OutputFormat.all in format:
        formats_to_export = [f for f in OutputFormat if f != OutputFormat.all]
    else:
        formats_to_export = format
    
    # Export each project in each format
    with Progress() as progress:
        task = progress.add_task("[green]Exporting projects...", total=len(projects))
        
        for project in projects:
            try:
                progress.update(task, description=f"[green]Exporting project: {project}...")
                
                # Create project directory
                project_dir = output_path / project
                project_dir.mkdir(parents=True, exist_ok=True)
                
                # Load the story
                story = story_persistence.load_story(project)
                
                # Get content
                content = story.content
                
                # Add metadata if requested
                if include_metadata:
                    metadata = story.metadata
                    metadata_text = f"""
# Story Metadata

- **Title**: {metadata.title}
- **Genre**: {metadata.genre}
- **Chapters**: {metadata.chapter_count}
- **Word Count**: {metadata.word_count:,}
- **Generated with**: {metadata.model if hasattr(metadata, "model") else "Unknown model"}
"""
                    content = metadata_text + "\n\n" + content
                
                # Export to each format
                for fmt in formats_to_export:
                    try:
                        output_file = project_dir / f"{project}.{fmt}"
                        
                        # Get the exporter function
                        exporter = get_exporter_for_format(fmt)
                        if not exporter:
                            console.print(f"[yellow]No exporter available for format: {fmt}[/yellow]")
                            continue
                        
                        # Export the content
                        exporter(content, str(output_file))
                        
                    except ImportError:
                        # Skip formats with missing dependencies in bulk mode
                        continue
                    except Exception as e:
                        console.print(f"[yellow]Error exporting {project} to {fmt}: {str(e)}[/yellow]")
                        continue
                
            except Exception as e:
                console.print(f"[yellow]Error processing project '{project}': {str(e)}[/yellow]")
                continue
            finally:
                progress.update(task, advance=1)
    
    console.print(f"[bold green]Bulk export complete![/bold green] Files saved to: {output_path}")


class ExportCommand(BaseCommand):
    """Command to export stories to various formats"""
    
    name = "export"
    help = "Export stories to various formats"
    
    @classmethod
    def _run_impl(
        cls,
        project: str = typer.Argument(..., help="Project name to export"),
        format: str = typer.Option(
            "markdown", "--format", "-f", 
            help="Output format to export to (plain, md, html, pdf, docx, epub, all)"
        ),
        output_file: Optional[str] = typer.Option(
            None, "--output", "-o", 
            help="Output file path (defaults to ./exports/PROJECT_NAME.FORMAT)"
        ),
        include_metadata: bool = typer.Option(
            False, "--metadata", "-m",
            help="Include story metadata in exported file"
        ),
    ):
        """Export a story to a specific format."""
        # Initialize story persistence to access stories
        story_dir = os.getenv("OUTPUT_DIR", "./output")
        story_persistence = StoryPersistence(story_dir)
        
        # Set up output file
        if not output_file:
            # Create exports directory if it doesn't exist
            exports_dir = Path("./exports")
            exports_dir.mkdir(parents=True, exist_ok=True)
            output_file = f"./exports/{project}.{format}"
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(os.path.abspath(output_file)), exist_ok=True)
        
        # Load the story
        try:
            story = story_persistence.load_story(project)
            cls.success(f"Loaded project: {project}")
        except Exception as e:
            cls.error(f"Error loading project '{project}': {str(e)}")
            return
        
        # Get content
        content = story.content
        
        # Add metadata if requested
        if include_metadata:
            metadata = story.metadata
            metadata_text = f"""
# Story Metadata

- **Title**: {metadata.title}
- **Genre**: {metadata.genre}
- **Chapters**: {metadata.chapter_count}
- **Word Count**: {metadata.word_count:,}
- **Generated with**: {metadata.model if hasattr(metadata, "model") else "Unknown model"}
"""
            content = metadata_text + "\n\n" + content
        
        # Export based on format
        try:
            if format.lower() == "all":
                # Export to all available formats
                formats = ["plain", "md", "html", "pdf", "docx", "epub"]
                output_base = os.path.splitext(output_file)[0]
                
                for fmt in formats:
                    try:
                        fmt_output = f"{output_base}.{fmt}"
                        exporter = get_exporter_for_format(fmt)
                        if exporter:
                            exporter(content, fmt_output)
                            cls.info(f"Exported to {fmt}: {fmt_output}")
                    except ImportError:
                        cls.warning(f"Missing dependency for {fmt} format")
                    except Exception as e:
                        cls.warning(f"Error exporting to {fmt}: {str(e)}")
                
                cls.success("Export complete!")
            else:
                # Export to a single format
                exporter = get_exporter_for_format(format.lower())
                if not exporter:
                    cls.error(f"Unsupported format: {format}")
                    return
                
                result_path = exporter(content, output_file)
                cls.success(f"Exported to {format}: {result_path}")
        
        except ImportError as e:
            cls.error(f"Missing dependency: {e}")
        except Exception as e:
            cls.error(f"Export error: {str(e)}") 