"""
DOCX exporter for Pulp Fiction Generator.
"""

from .base import BaseExporter
from .exceptions import ExporterFileError, ExporterDependencyError
from ...utils.errors import logger


class DocxExporter(BaseExporter):
    """Exporter for DOCX format."""
    
    @staticmethod
    def check_dependencies() -> bool:
        """
        Check if python-docx is installed.
        
        Returns:
            bool: True if dependencies are available, False otherwise
        """
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def export(self, content: str, output_path: str) -> str:
        """
        Export content as DOCX.
        
        Args:
            content: The story content to export
            output_path: Path where the exported file should be saved
            
        Returns:
            The path to the exported file
            
        Raises:
            ExporterDependencyError: If python-docx is not installed
            ExporterFileError: If unable to write to the output file
        """
        try:
            # Import dependencies
            try:
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                raise ExporterDependencyError(
                    "python-docx package is required for DOCX export. "
                    "Install with: pip install python-docx"
                )
                
            document = Document()
            
            # Extract title from the content (first line)
            lines = content.strip().split("\n")
            title = self.extract_title(content) or "Generated Story"
            
            # Add title
            document.add_heading(title, 0)
            
            # Process content by paragraphs
            current_para = ""
            in_heading = False
            
            for line in lines[1:]:
                # Skip the title line
                if line.strip() == "" and current_para:
                    # End of paragraph
                    document.add_paragraph(current_para)
                    current_para = ""
                elif line.startswith('#'):
                    # This is a heading
                    if current_para:
                        document.add_paragraph(current_para)
                        current_para = ""
                    
                    # Add the heading (count # to determine level)
                    level = 0
                    heading_text = line.strip()
                    while heading_text.startswith("#"):
                        level += 1
                        heading_text = heading_text[1:]
                    
                    # Cap at level 9 (docx only supports up to 9)
                    level = min(level, 9)
                    document.add_heading(heading_text.strip(), level)
                else:
                    # Regular paragraph text
                    if current_para:
                        current_para += " " + line.strip()
                    else:
                        current_para = line.strip()
            
            # Add the last paragraph if any
            if current_para:
                document.add_paragraph(current_para)
            
            # Save the document
            document.save(output_path)
            return output_path
            
        except ExporterDependencyError:
            # Re-raise dependency errors
            raise
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise ExporterFileError(f"Error exporting to DOCX: {e}") from e 